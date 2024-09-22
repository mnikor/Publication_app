
import os
import re
import json
import logging
from typing import Dict, Any, Optional, List, Tuple
from functools import lru_cache
from io import BytesIO
import textstat  # Add this import for readability calculations
from collections import Counter
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import PyPDF2
import docx
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
import tempfile
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
from lifelines import KaplanMeierFitter
import seaborn as sns
from matplotlib.colors import LinearSegmentedColormap
import markdown2
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle

# Set up logging
logging.basicConfig(level=logging.DEBUG)


# Initialize OpenAI with the API key from Streamlit secrets
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

# Optional: Verify that the API key is loaded (for debugging purposes only; remove in production)
# st.write(f"OpenAI API Key Loaded: {'Yes' if openai.api_key else 'No'}")
  # Replace with your actual API key

# Update the PUBLICATION_TYPES dictionary
PUBLICATION_TYPES = {
    "Congress Abstract": {
        "max_characters": 2000,  # Changed from max_words to max_characters
        "font_sizes": {
            "title": 14,
            "body": 11
        },
        "structure": [
            "Title",
            "Authors",
            "Affiliations",
            "Background",
            "Methods",
            "Results",
            "Conclusions",
            "Funding",
            "Keywords"
        ]
    },
    "Manuscript": {
        "max_words": 3500,
        "font_sizes": {
            "title": 16,
            "heading": 14,
            "subheading": 12,
            "body": 11
        },
        "structure": [
            "Title",
            "Authors",
            "Affiliations",
            "Abstract",
            "Introduction",
            "Methods",
            "Results",
            "Discussion",
            "Conclusion",
            "Acknowledgements",
            "References",
            "Tables and Figures"
        ]
    },
    "Poster": {
        "max_characters": 10000,
        "font_sizes": {
            "title": 60,
            "heading": 36,
            "subheading": 24,
            "body": 18
        },
        "structure": [
            "Title",
            "Authors",
            "Affiliations",
            "Introduction",
            "Methods",
            "Results",
            "Conclusion",
            "References",
            "Acknowledgements"
        ]
    },
    "Plain Language Summary": {
        "max_words": 750,
        "min_words": 200,
        "font_sizes": {
            "title": 16,
            "body": 12
        },
        "structure": [
            "Title",
            "Key Points",
            "Background",
            "What was the study about?",
            "How was the study done?",
            "What were the results?",
            "What do the results mean for patients?",
            "What's next?",
            "Disclosures",
            "Review"
        ]
    }
}

# Update the ANALYSIS_TYPES dictionary to include all types
ANALYSIS_TYPES = {
    "Primary Efficacy Analysis": {
        "max_words": 2000,
        "font_sizes": {
            "title": 16,
            "heading": 14,
            "body": 12
        },
        "structure": [
            "Title",
            "Authors",
            "Affiliations",
            "Abstract",
            "Introduction",
            "Methods",
            "Primary Results",
            "Discussion",
            "Conclusion",
            "References",
            "Keywords"
        ]
    },
    "Safety Analysis": {
        "max_words": 2000,
        "font_sizes": {
            "title": 16,
            "heading": 14,
            "body": 12
        },
        "structure": [
            "Title",
            "Authors",
            "Affiliations",
            "Abstract",
            "Introduction",
            "Methods",
            "Safety Results",
            "Discussion",
            "Conclusion",
            "References",
            "Keywords"
        ]
    },
    "Pharmacokinetic (PK) Analysis": {
        "max_words": 2000,
        "font_sizes": {
            "title": 16,
            "heading": 14,
            "body": 12
        },
        "structure": [
            "Title",
            "Authors",
            "Affiliations",
            "Abstract",
            "Introduction",
            "Methods",
            "PK Results",
            "Discussion",
            "Conclusion",
            "References",
            "Keywords"
        ]
    },
    "Subgroup Analysis": {
        "max_words": 2000,
        "font_sizes": {
            "title": 16,
            "heading": 14,
            "body": 12
        },
        "structure": [
            "Title",
            "Authors",
            "Affiliations",
            "Abstract",
            "Introduction",
            "Methods",
            "Subgroup Results",
            "Discussion",
            "Conclusion",
            "References",
            "Keywords"
        ]
    },
    "Interim Analysis": {
        "max_words": 1500,
        "font_sizes": {
            "title": 16,
            "heading": 14,
            "body": 12
        },
        "structure": [
            "Title",
            "Authors",
            "Affiliations",
            "Background",
            "Study Design",
            "Interim Objectives",
            "Methods",
            "Interim Results",
            "Discussion",
            "Conclusion",
            "Keywords"
        ]
    },
    "Post-hoc Analysis": {
        "max_words": 2000,
        "font_sizes": {
            "title": 16,
            "heading": 14,
            "body": 12
        },
        "structure": [
            "Title",
            "Authors",
            "Affiliations",
            "Abstract",
            "Introduction",
            "Methods",
            "Post-hoc Results",
            "Discussion",
            "Conclusion",
            "References",
            "Keywords"
        ]
    },
    "Trial in Progress": {
        "max_words": 500,
        "font_sizes": {
            "title": 16,
            "heading": 14,
            "body": 12
        },
        "structure": [
            "Title",
            "Authors",
            "Affiliations",
            "Background",
            "Study Design",
            "Objectives",
            "Methods",
            "Current Status",
            "Conclusion",
            "Keywords"
        ]
    },
    "Baseline Characteristics": {
        "max_words": 1000,
        "font_sizes": {
            "title": 16,
            "heading": 14,
            "body": 12
        },
        "structure": [
            "Title",
            "Authors",
            "Affiliations",
            "Introduction",
            "Methods",
            "Baseline Characteristics",
            "Results",
            "Conclusion",
            "Keywords"
        ]
    }
}

# The ANALYSIS_SOURCE_RECOMMENDATIONS dictionary is already correct and complete
# as it includes all the analysis types from the table you provided.

# Cache the document generation to improve performance (unchanged)
@lru_cache(maxsize=128)
def generate_document_cached(publication_type: str, analysis_type: str, user_input: str, additional_instructions: str) -> Optional[Dict[str, Any]]:
    return generate_document(publication_type, analysis_type, user_input, additional_instructions)

def get_section_requirements(publication_type: str) -> str:
    if publication_type == "Congress Abstract":
        return """
           - Title: Concise and reflective of the abstract content. Do not include study results or conclusions. Use significant words that reflect the content. Do not use commercial names; use generic names in lower case.
           - Authors: Full names of all authors, typically without credentials.
           - Affiliations: Brief institutional affiliations including institution and location.
           - Background: A short introduction indicating the rationale of the study (1-2 sentences).
           - Methods: A brief description of pertinent methodological procedures (2-3 sentences).
           - Results: A summary of the key findings with essential statistical results (3-4 sentences).
           - Conclusions: A statement of the main conclusions (1-2 sentences).
           - Funding: Provide the name(s) of the legal entity/entities responsible for the study and the organisation(s) providing funding.
           - Keywords: 3-5 relevant terms.
           - General guidelines:
             - Maximum 2,000 characters (excluding spaces) for the entire abstract.
             - Define abbreviations at first mention.
             - For commercial names, use the format: generic (Commercial®).
             - Identify complex chemotherapeutic regimens clearly.
        """
    elif publication_type == "Plain Language Summary":
        return """
           - Title: Simple, clear, and reflective of the main message. (10-15 words)
           - Summary: Concise explanation of the study and its findings in patient-friendly language. (200-750 words total)
             Structure the summary as follows:
             1. Key Points: 3-5 bullet points summarizing the most important takeaways.
             2. Background: Brief context about the condition and why the study was done. (2-3 sentences)
             3. What was the study about?: Clear statement of the study's purpose. (1-2 sentences)
             4. How was the study done?: Simple description of the study methods, avoiding technical details. (3-4 sentences)
             5. What were the results?: Key findings in plain language, focusing on what's most relevant to patients. (3-5 sentences)
             6. What do the results mean for patients?: Practical implications of the findings for patient care or decision-making. (2-3 sentences)
             7. What's next?: Brief mention of any study limitations or ongoing research questions. (1-2 sentences)
           - Writing Guidelines:
             - Use a 6th to 8th-grade reading level (Flesch-Kincaid grade level 6-8).
             - Keep sentences short (15-20 words on average) and paragraphs brief (3-5 sentences).
             - Use active voice and present tense where possible.
             - Avoid jargon, acronyms, and abbreviations. If used, explain them immediately.
             - Use common, everyday words instead of technical terms.
             - Address questions patients are likely to have.
           - Visual Elements: If including any charts or graphics, ensure they are simple and clearly labeled.
           - Disclosures: Include funding sources and any potential conflicts of interest.
           - Review: Mention that the summary has been reviewed by both a medical expert and a patient advocate (if applicable).
        """
   elif publication_type == "Manuscript":
       return """
    - Title:
      • Concise, informative, and reflective of the study's main focus (typically 10-20 words)
      • Include key elements such as study design, population, and primary outcome
      • Avoid abbreviations and declarative statements of results

    - Authors and Affiliations:
      • Full names, highest academic degrees, and institutional affiliations of all authors
      • Clearly indicate the corresponding author and provide their full contact details
      • Include any study group name if applicable

    - Abstract:
      • Structured summary (250-300 words) including:
        - Background: Context and rationale for the study
        - Objective: Specific aims or hypotheses
        - Methods: Basic study design, participants, interventions, and main outcome measures
        - Results: Key findings with effect sizes and precision (e.g., 95% confidence interval)
        - Conclusions: Main interpretations and implications of the results
      • Provide the trial registration number at the end

    - Introduction (typically 3-5 paragraphs):
      • Provide a compelling rationale for the study based on existing literature
      • Clearly state the knowledge gap this study addresses
      • Define the specific objective(s) or hypothesis(es)
      • Briefly outline the approach

    - Methods:
      • Study Design:
        - Describe the type of study (e.g., randomized controlled trial, cohort study)
        - Specify any changes to methods after trial commencement, with reasons
      • Participants:
        - Detail eligibility criteria, settings, and locations where data were collected
        - Describe recruitment procedures and any pre-randomization assessments
      • Interventions:
        - Provide precise details of the interventions for each group
        - Include how and when they were actually administered
      • Outcomes:
        - Define all primary and secondary outcome measures
        - Specify their measurement methods and timing
      • Sample Size:
        - Explain how sample size was determined
        - Describe any interim analyses and stopping guidelines
      • Randomization:
        - Detail the method used to generate the random allocation sequence
        - Describe the mechanism used to implement the allocation sequence
        - Specify who generated the sequence, enrolled participants, and assigned them to interventions
      • Blinding:
        - Describe who was blinded after assignment to interventions
      • Statistical Methods:
        - Describe methods used to compare groups for primary and secondary outcomes
        - Explain methods for any additional analyses (e.g., subgroup or adjusted analyses)

    - Results:
      • Participant Flow:
        - Provide a diagram showing participant flow through each stage of the study
        - For each group, report the numbers randomly assigned, receiving intended treatment, and analyzed for the primary outcome
      • Recruitment:
        - State the dates defining the periods of recruitment and follow-up
      • Baseline Data:
        - Provide a table showing baseline demographic and clinical characteristics for each group
      • Numbers Analyzed:
        - For each group, report numbers of participants analyzed and whether analysis was by original assigned groups
      • Outcomes and Estimation:
        - For each primary and secondary outcome, provide results for each group, the estimated effect size, and its precision (e.g., 95% confidence interval)
        - For binary outcomes, present both absolute and relative effect sizes
      • Ancillary Analyses:
        - Report results of any other analyses performed, including subgroup analyses and adjusted analyses, distinguishing pre-specified from exploratory
      • Harms:
        - Report all important harms or unintended effects in each group

    - Discussion:
      • Key Findings:
        - Summarize key results with reference to study objectives
      • Limitations:
        - Discuss limitations of the study, addressing sources of potential bias or imprecision
      • Interpretation:
        - Provide a general interpretation of the results in the context of current evidence
      • Generalizability:
        - Discuss the external validity (generalizability) of the trial findings

    - Conclusion:
      • Summarize the main findings
      • State their importance and implications for practice or future research

    - Acknowledgements:
      • Recognize all persons who contributed but do not meet authorship criteria
      • Specify their contributions
      • List all sources of funding and the role of funders in the study

    - References:
      • Cite 30-50 relevant and recent sources for original research
      • Follow journal-specific formatting guidelines
      • Ensure all in-text citations are included and vice versa

    - Tables and Figures:
      • Include 5-7 essential visual representations of data
      • Provide clear, self-explanatory titles and legends
      • Ensure all abbreviations are explained in footnotes

    General Guidelines:
    - Maximum 3,500 words (excluding abstract, references, tables, and figures)
    - Use clear, precise language appropriate for a scientific audience, prioritizing thoroughness over brevity
    - Provide comprehensive explanations and detailed descriptions for each section
    - Include all relevant information, data, and context necessary for a complete understanding of the study
    - Elaborate on complex concepts, methodologies, and findings
    - Use specific examples, detailed explanations, and thorough justifications
    - Aim for depth and breadth in coverage while maintaining logical flow
    - Use topic sentences followed by supporting details to build comprehensive paragraphs
    - Incorporate relevant background information and contextual details
    - Provide sufficient detail in methods to allow for potential replication
    - Use precise terminology and define specialized terms or acronyms upon first use
    - Quantify findings whenever possible, providing exact figures, percentages, and statistical analyses
    - Address potential questions or areas of ambiguity proactively
    - Ensure each section is self-contained and comprehensive, while connecting logically to other sections
    - Follow standard scientific writing conventions (e.g., IMRAD structure)
    - Use past tense for completed actions and present tense for known facts and conclusions
    - If in doubt about including information, err on the side of inclusion
    - Aim for a thorough, authoritative account of the study that leaves no important questions unanswered
    """
else:
        return ""

def calculate_flesch_kincaid_grade(text: str) -> float:
    return textstat.flesch_kincaid_grade(text)

def assess_content_quality(content: str, publication_type: str, analysis_type: str) -> Dict[str, Any]:
    assessment = {}

    # 1. Readability Scores
    assessment["readability"] = {
        "flesch_kincaid_grade": textstat.flesch_kincaid_grade(content),
        "flesch_reading_ease": textstat.flesch_reading_ease(content),
        "smog_index": textstat.smog_index(content),
    }

    # 2. Word Count and Section Balance
    sections = re.split(r'\n##\s+', content)
    section_word_counts = {section.split('\n')[0]: len(section.split()) for section in sections if section.strip()}
    assessment["word_counts"] = section_word_counts
    total_words = sum(section_word_counts.values())
    assessment["total_words"] = total_words

    # 3. Character Count (for Congress Abstract)
    if publication_type == "Congress Abstract":
        assessment["total_characters"] = len(content)

    # 4. Keyword Density
    words = re.findall(r'\w+', content.lower())
    word_freq = Counter(words)
    total_words = len(words)
    keyword_density = {word: count/total_words for word, count in word_freq.most_common(10)}
    assessment["keyword_density"] = keyword_density

    # 5. Citation Check
    citation_pattern = r'\(\w+\s+et\s+al\.,\s+\d{4}\)|\[\d+\]'
    citations = re.findall(citation_pattern, content)
    assessment["citation_count"] = len(citations)

    # 6. AI-powered Content Evaluation
    openai_api_key = st.secrets['openai']["OPENAI_API_KEY"]
    
    # Truncate content to approximately 128000 tokens (assuming 4 characters per token)
    max_input_chars = 128000 * 4
    truncated_content = content[:max_input_chars]
    
    if publication_type == "Plain Language Summary":
        prompt = f"""
        Evaluate the following Plain Language Summary for a {analysis_type}. 
        Provide a comprehensive assessment of its quality, focusing on:
        1. Clarity and simplicity of language (aim for 6th to 8th-grade reading level)
        2. Avoidance of jargon and technical terms
        3. Logical flow and organization of information
        4. Relevance to patient audience
        5. Inclusion of key sections (Background, Purpose, Methods, Results, Implications)
        6. Use of everyday examples or analogies to explain complex concepts
        7. Appropriate length (200-750 words)

        Highlight any areas that need improvement and suggest specific enhancements.

        Content:
        {truncated_content}

        Evaluation:
        """
    else:
        prompt = f"""
        Evaluate the following {publication_type} content for a {analysis_type}. 
        Provide a comprehensive assessment of its quality, coherence, and adherence to scientific writing standards.
        Highlight any areas that need improvement and suggest specific enhancements.

        Content:
        {truncated_content}

        Evaluation:
        """
def assess_content_quality(content: str, publication_type: str, analysis_type: str) -> Dict[str, Any]:
    assessment = {}
    try:
        # 1. Readability Scores
        assessment["readability"] = {
            "flesch_kincaid_grade": textstat.flesch_kincaid_grade(content),
            "flesch_reading_ease": textstat.flesch_reading_ease(content),
            "smog_index": textstat.smog_index(content),
        }

        # 2. Word Count and Section Balance
        try:
            sections = re.split(r'\n##\s+', content)
            section_word_counts = {section.split('\n')[0]: len(section.split()) for section in sections if section.strip()}
            assessment["word_counts"] = section_word_counts
            total_words = sum(section_word_counts.values())
            assessment["total_words"] = total_words
        except Exception as e:
            logging.error(f"Error calculating word counts: {str(e)}")
            assessment["word_counts"] = {}
            assessment["total_words"] = len(content.split())  # Fallback to simple word count

        # 3. Character Count (for Congress Abstract)
        if publication_type == "Congress Abstract":
            assessment["total_characters"] = len(content)

        # 4. Keyword Density
        words = re.findall(r'\w+', content.lower())
        word_freq = Counter(words)
        total_words = len(words)
        keyword_density = {word: count/total_words for word, count in word_freq.most_common(10)}
        assessment["keyword_density"] = keyword_density

        # 5. Citation Check
        citation_pattern = r'\(\w+\s+et\s+al\.,\s+\d{4}\)|\[\d+\]'
        citations = re.findall(citation_pattern, content)
        assessment["citation_count"] = len(citations)

        # 6. AI-powered Content Evaluation
        try:
            max_input_chars = 16000 * 4
            truncated_content = content[:max_input_chars]
            
            prompt = f"""
            Evaluate the following {publication_type} content for a {analysis_type}. 
            Provide a comprehensive assessment of its quality, coherence, and adherence to scientific writing standards.
            Highlight any areas that need improvement and suggest specific enhancements.

            Content:
            {truncated_content}

            Evaluation:
            """
            
            response = client.chat.completions.create(
                model="gpt-4o-2024-08-06",
                messages=[
                    {"role": "system", "content": "You are a scientific writing expert."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1000
            )
            
            assessment["ai_evaluation"] = response.choices[0].message.content
        except Exception as e:
            logging.error(f"Error in AI evaluation: {str(e)}")
            assessment["ai_evaluation"] = "AI evaluation failed due to an error."

    except Exception as e:
        logging.exception("Error in assess_content_quality:")
        assessment["error"] = str(e)

    return assessment
    
    assessment["ai_evaluation"] = response.choices[0].message.content

    return assessment

def extract_tabular_data(text: str) -> str:
    """
    Extracts potential tabular data from the input text.
    This is a simple implementation and might need to be enhanced based on the actual format of your source documents.
    """
    # Look for patterns that might indicate tabular data
    table_pattern = r'(\|.*\|[\n\r])+\|.*\|'
    tables = re.findall(table_pattern, text)
    
    # Join all found tables
    extracted_data = "\n\n".join(tables)
    
    return extracted_data if extracted_data else "No tabular data found in the source document."

def generate_document(publication_type: str, analysis_type: str, user_input: str, additional_instructions: str) -> Optional[Dict[str, Any]]:
    try:
        pub_type_info = PUBLICATION_TYPES[publication_type]
        analysis_type_info = ANALYSIS_TYPES[analysis_type]
        
        max_length_pub = pub_type_info.get("max_words", pub_type_info.get("max_characters", ""))
        length_type_pub = "words" if "max_words" in pub_type_info else "characters"
        
        max_length_analysis = analysis_type_info.get("max_words", analysis_type_info.get("max_characters", ""))
        length_type_analysis = "words" if "max_words" in analysis_type_info else "characters"
        
        font_sizes = {**pub_type_info["font_sizes"], **analysis_type_info["font_sizes"]}
        font_size_info = ", ".join([f"{k.capitalize()}: {v}pt" for k, v in font_sizes.items()])

        structure = list(dict.fromkeys(pub_type_info["structure"] + analysis_type_info["structure"]))
        structure_info = "\n".join([f"- {section}" for section in structure])

        extracted_data = extract_tabular_data(user_input)

        if publication_type == "Plain Language Summary":
            prompt = f"""
            You are a professional scientific medical writing assistant specializing in transforming Clinical Study Reports (CSRs) and other source documents into various publication types.

            You are tasked with generating a comprehensive Plain Language Summary that combines the structure and guidelines of the following:

            **Publication Type:** {publication_type}
            **Analysis Type:** {analysis_type}

            ### **Guidelines:**

            **Target Reading Level:**
            - Write the summary at a 6th to 8th-grade reading level.
            - Aim for short sentences averaging 15 words or fewer.
            - Use simple sentence structures; avoid complex or compound sentences.

            **Language and Style:**
            - Use common, everyday words instead of medical jargon.
            - If medical terms are necessary, explain them in simple language.
            - Write in active voice and present tense where appropriate.
            - Engage the reader by addressing them directly when suitable.

            **Structure and Content:**
            - **Title:** Simple and clear, reflecting the main message (10-15 words).
            - **Key Points:** 3-5 bullet points summarizing the most important takeaways.
            - **Background:** Brief context about the condition and why the study was done (2-3 sentences).
            - **What Was the Study About?:** Clear statement of the study's purpose (1-2 sentences).
            - **How Was the Study Done?:** Simple description of the study methods, avoiding technical details (2-3 sentences).
            - **What Were the Results?:** Key findings in plain language, focusing on what's most relevant to patients (3-4 sentences).
            - **What Do the Results Mean for Patients?:** Practical implications for patient care or decision-making (2-3 sentences).
            - **What's Next?:** Mention any study limitations or ongoing research (1-2 sentences).
            - **Disclosures:** Include funding sources and any potential conflicts of interest.
            - **Review Statement:** State that the summary was reviewed by a medical expert and a patient advocate (if applicable).

            **Acronyms and Abbreviations:**
            - Spell out acronyms upon first use and provide a simple explanation if necessary.

            **Visual Aids:**
            - If helpful, include simple visual elements to explain key concepts.
            - Ensure visuals are clearly labeled and easy to understand.

            **Writing Tips:**
            - Keep paragraphs brief (3-5 sentences).
            - Use bullet points or numbered lists where appropriate.
            - Address common questions patients might have.
            - Avoid unnecessary words or filler content.

            **Final Review:**
            - Before finalizing, read the summary aloud to ensure it flows naturally.
            - Verify that the FKGL is between 6 and 8 using readability assessment tools.
            - Make adjustments to sentence length and word choice as needed to achieve the target reading level.

            Input:
            {user_input}

            Additional Instructions:
            {additional_instructions}
            """
        elif publication_type == "Congress Abstract":
            prompt = f"""
            You are a professional scientific medical writing assistant specializing in transforming Clinical Study Reports (CSRs) and other source documents into various publication types.

            You are tasked with generating a scientific congress abstract following these guidelines:

            **Publication Type:** {publication_type}
            **Analysis Type:** {analysis_type}

            ### **Guidelines:**

            1. **Structure:**
               Create an abstract with the following four sections:
               a) Background: Provide a brief introduction explaining the study's rationale.
               b) Methods: Describe the key methodological procedures concisely.
               c) Results: Summarize the main findings of the research.
               d) Conclusions: State the primary conclusions drawn from the study.

            2. **Title:**
               - Craft a title that reflects the abstract's content using significant words.
               - Do not include study results or conclusions in the title.
               - Avoid using commercial names in the title.

            3. **Content Guidelines:**
               - Use generic names for compounds in lower case.
               - If including commercial names in the text, use the ® symbol and place them in brackets after the generic name, e.g., "generic (Commercial®)".
               - Provide the name(s) of the legal entity/entities responsible for the study's governance, coordination, and execution.
               - Include the name(s) of organizations providing funding.

            4. **Abbreviations:**
               - Define all abbreviations upon first use.
               - Spell out terms in full at first mention, followed by the abbreviation in parentheses.
               - Take extra care to identify complex chemotherapeutic regimens clearly.

            5. **Length:**
               - Limit the abstract to 2,000 characters, excluding spaces.

            6. **Additional Notes:**
               - Ensure all information is accurate and reflects the study correctly.
               - Maintain a professional and scientific tone throughout the abstract.
               - Focus on presenting the most crucial and impactful aspects of the study within the limited space.

            Input:
            {user_input}

            Additional Instructions:
            {additional_instructions}
            """
        elif publication_type == "Manuscript":
            prompt = f"""
            You are a professional scientific medical writing assistant. Generate a comprehensive manuscript for a clinical trial, adhering to international reporting standards (e.g., CONSORT). Ensure ALL of the following sections and details are included:

            1. Title: Concise, informative title reflecting the study's main focus.

            2. Abstract: Structured summary (250-300 words) including:
               - Background
               - Methods (design, participants, interventions, main outcomes)
               - Results (number randomized, primary and key secondary outcomes)
               - Conclusions
               - Trial Registration: NCT number

            3. Introduction:
               - Background: Explain the scientific context and rationale
               - Objectives: State specific objectives and hypotheses

 4. Methods:
   - Trial Design:
     • Describe the type of trial (e.g., parallel group, crossover, factorial) and its framework (e.g., superiority, equivalence, non-inferiority)
     • Specify the allocation ratio
     • Detail any important changes to methods after trial commencement (e.g., eligibility criteria), with reasons
   - Participants:
     • Provide detailed eligibility criteria for participants, including age range, gender, diagnosis criteria, and any specific inclusion/exclusion criteria
     • Describe the settings and locations where the data were collected (e.g., number and type of centers)
     • Specify the duration of the study, including recruitment period and follow-up
   - Interventions:
     • For each group, provide precise details of the interventions:
       - The drug name (generic and trade), manufacturer, dose, route of administration, and regimen
       - Any washout periods and run-in phases
       - Procedures for monitoring participant compliance
     • Describe any concomitant treatments allowed or prohibited
   - Outcomes:
     • Clearly define all primary and secondary outcome measures, including how and when they were assessed
     • Specify any changes to trial outcomes after the trial commenced, with reasons
     • For each outcome measure, provide details on its validity and reliability in the study population
   - Sample Size:
     • Explain how sample size was determined
     • Provide details of any interim analyses and stopping guidelines, if applicable
   - Randomization:
     • Sequence Generation: Describe the method used to generate the random allocation sequence (e.g., computer-generated random numbers)
     • Allocation Concealment: Detail the mechanism used to implement the random allocation sequence (e.g., sequentially numbered containers), describing any steps taken to conceal the sequence until interventions were assigned
     • Implementation: Specify who generated the random allocation sequence, who enrolled participants, and who assigned participants to interventions
   - Blinding:
     • Describe who was blinded after assignment to interventions (e.g., participants, care providers, outcome assessors)
     • If relevant, describe the similarity of interventions (e.g., appearance, taste, smell)
     • Provide details on how the success of blinding was evaluated
   - Statistical Methods:
     • Describe statistical methods used to compare groups for primary and secondary outcomes, including:
       - Handling of missing data (e.g., complete-case analysis, multiple imputation)
       - Any adjustments for multiple comparisons
       - Analysis population (e.g., intention-to-treat, per-protocol)
     • Explain methods for additional analyses, such as subgroup analyses, adjusted analyses, and sensitivity analyses
     • Specify the statistical software used
   - Ethical Considerations:
     • Provide details of ethical approval, including the name of the ethics committee(s) and the approval number(s)
     • Describe the process for obtaining informed consent from participants
     • Mention adherence to specific ethical guidelines (e.g., Declaration of Helsinki)
     • Detail any data monitoring and safety procedures

5. Results:
   - Participant Flow:
     • Provide a detailed flow diagram of participant progress through the study, including:
       - Number of participants screened, eligible, consented, randomized, received intended treatment, and analyzed for the primary outcome
       - Reasons for non-participation at each stage
     • For each group, describe losses and exclusions after randomization, with reasons
   - Recruitment:
     • State the specific dates defining the periods of recruitment and follow-up
     • Explain why the trial ended or was stopped, if applicable
   - Baseline Data:
     • Present a detailed table showing baseline demographic and clinical characteristics for each group
     • Include all important prognostic variables and potential confounding factors
   - Numbers Analyzed:
     • For each group, report the number of participants included in each analysis
     • Specify whether the analysis was by original assigned groups
     • Provide reasons for any protocol deviations
   - Outcomes and Estimation:
     • For each primary and secondary outcome:
       - Present results for each group
       - Provide the estimated effect size and its precision (e.g., 95% confidence interval)
       - For binary outcomes, present both absolute and relative effect sizes
     • Include results of any other analyses performed, including subgroup analyses and adjusted analyses, distinguishing pre-specified from exploratory analyses
     • Present results in a clear, organized manner, using tables and figures where appropriate
   - Ancillary Analyses:
     • Report the results of any other analyses performed, including:
       - Subgroup analyses
       - Adjusted analyses
       - Sensitivity analyses
     • Clearly distinguish pre-specified from exploratory analyses
   - Harms:
     • Provide a comprehensive summary of all important harms or unintended effects in each group
     • Include:
       - Absolute risks of important adverse events for each group
       - Serious adverse events
       - Withdrawals due to adverse events
       - Most common adverse events (e.g., those occurring in >5% of participants)
     • Describe any differences in adverse events between treatment groups
     • Present data on the timing, severity, and duration of adverse events, if available

Ensure that each of these sections is thoroughly addressed, providing specific details and data from the study wherever possible. Use tables, figures, and statistical analyses as appropriate to present the information clearly and comprehensively.
            6. Discussion:
               - Limitations: Trial limitations, addressing sources of potential bias, imprecision, and, if relevant, multiplicity of analyses
               - Generalizability: External validity of the trial findings
               - Interpretation: Interpretation consistent with results, balancing benefits and harms, and considering other relevant evidence

            7. Other Information:
               - Registration: Registration number and name of trial registry
               - Protocol: Where the full trial protocol can be accessed, if available
               - Funding: Sources of funding and other support, role of funders

            Ensure that all sections are comprehensive and adhere to the guidelines provided. Use the following input to generate the manuscript:

            Analysis Type: {analysis_type}
            Maximum Length: {max_length_pub} {length_type_pub}
            Font Sizes: {font_size_info}
            Structure: {structure_info}
            Extracted Data: {extracted_data}
            Input: {user_input}
            Additional Instructions: {additional_instructions}

            Generate the complete manuscript, ensuring all required sections and details are included. Adhere to the specified maximum length and structure guidelines.
            """
        else:
            prompt = f"""
            You are a professional scientific medical writing assistant specializing in transforming Clinical Study Reports (CSRs) and other source documents into various publication types.

            You are tasked with generating a comprehensive document that combines the structure and guidelines of the following:

            **Publication Type:** {publication_type}
            **Analysis Type:** {analysis_type}

            ### **Guidelines:**

            1. **Document Length:**
               - **Publication:** Maximum {max_length_pub} {length_type_pub}.
               - **Analysis:** Maximum {max_length_analysis} {length_type_analysis}.

            2. **Font Sizes:**
               - {font_size_info}

            3. **Structure:**
               - The document should include all sections from both the publication type and analysis type. Ensure that each section is clearly marked using Markdown syntax (e.g., ## Title, ### Methods).
               - Provide detailed and comprehensive content for each section. Aim for at least 2-3 sentences per section, unless otherwise specified.

            4. **Content Generation:**
               - Use clear and concise language appropriate for a scientific publication.
               - If specific information is not provided in the input, use placeholder text or general statements that would be appropriate for the section.

            5. **Visualizations:**
               - Extract key numerical data from the input and suggest up to 2 relevant charts or visualizations.
               - For each chart, provide the following in JSON format, enclosed within triple backticks and specify the language as JSON:

            ```json
            {{
              "type": "Chart Type (e.g., Bar Chart, Line Chart)",
              "title": "Chart Title",
              "x_label": "X-axis Label",
              "y_label": "Y-axis Label",
              "data_series": ["Numerical Series1", "Numerical Series2", ...],
              "data": [
                {{"X-axis Value": ..., "Numerical Series1": ..., "Numerical Series2": ...}},
                ...
              ]
            }}
            ```

            After completing the publication and analysis content, provide a separate section titled "## Visualizations" containing all chart JSON data.

            6. **Tables:**
               - Include up to 5-7 essential tables that complement the text.
               - For each table:
                 - Provide a detailed title
                 - List column headers
                 - Use actual data from the source document if available. Here's the extracted tabular data:
                   {extracted_data}
                 - If actual data is not available or incomplete, provide placeholder data or ranges based on the study information
               - Use Markdown table syntax for creating tables.

            7. **Acknowledgement:**
               - ALWAYS include an Acknowledgement section at the end of the document with the following text:
                 "This [publication type] was created with the assistance of generative AI technology."

            Adherence to Guidelines:
            Strictly adhere to the format and guidelines for both the publication type and analysis type.
            Ensure that ALL sections specified in the combined structure are present and contain at least minimal content.

            Combined Structure:
            {structure_info}

            Input:
            {user_input}

            Additional Instructions:
            {additional_instructions}
            """

        response = client.chat.completions.create(
            model="gpt-4o-2024-08-06",
            messages=[
                {"role": "system", "content": "You are a professional scientific medical writing assistant specializing in transforming Clinical Study Reports (CSRs) and other source documents into various publication types."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=16000,
            temperature=0  # Ensures consistency
        )

        full_content = response.choices[0].message.content

        # Extract chart information
        charts = extract_chart_info(full_content)
        
        logging.debug(f"Extracted charts: {charts}")

        return {"content": full_content, "charts": charts}

    except Exception as e:
        logging.error(f"Error in generate_document: {str(e)}")
        return {"content": f"An error occurred while generating the document: {str(e)}", "charts": []}

def extract_chart_info(content: str) -> List[Dict[str, Any]]:
    """
    Extracts chart information from the '## Visualizations' section of the generated content.

    Parameters:
    - content (str): The full generated content containing chart JSON.

    Returns:
    - List[Dict[str, Any]]: A list of chart information dictionaries.
    """
    charts = []
    # Locate the '## Visualizations' section
    visualizations_match = re.search(r'##\s+Visualizations\s*([\s\S]*)', content, re.IGNORECASE)
    if visualizations_match:
        visualizations_content = visualizations_match.group(1)
        logging.debug(f"Found 'Visualizations' section:\n{visualizations_content}")
        # Find all JSON blocks within the 'Visualizations' section
        json_blocks = re.findall(r'```json\s*([\s\S]*?)```', visualizations_content, re.DOTALL | re.IGNORECASE)
        logging.debug(f"Found {len(json_blocks)} JSON blocks in 'Visualizations' section.")
        for idx, json_str in enumerate(json_blocks, 1):
            try:
                json_str = json_str.strip()
                chart_info = json.loads(json_str)
                # Validate required fields
                if not validate_chart_data(chart_info):
                    logging.warning(f"Chart {idx} is missing required fields or has invalid data.")
                    continue
                charts.append(chart_info)
                logging.debug(f"Chart {idx} extracted successfully.")
            except json.JSONDecodeError as e:
                logging.error(f"JSON decoding error in chart {idx}: {e} in block: {json_str}")
                continue
    else:
        logging.warning("No 'Visualizations' section found in the generated content.")
    return charts

def validate_chart_data(chart_info: Dict[str, Any]) -> bool:
    required_fields = {"type", "title", "x_label", "y_label", "data_series", "data"}
    if not required_fields.issubset(chart_info.keys()):
        logging.error(f"Chart info missing required fields: {chart_info}")
        return False
    
    # Validate data_series is a list of strings
    if not isinstance(chart_info['data_series'], list) or not all(isinstance(series, str) for series in chart_info['data_series']):
        logging.error("data_series must be a list of strings.")
        return False
    
    # Validate data is a non-empty list of dictionaries
    if not isinstance(chart_info['data'], list) or not chart_info['data']:
        logging.error("Chart data is not a non-empty list.")
        return False
    for entry in chart_info['data']:
        if not isinstance(entry, dict):
            logging.error(f"Data entry is not a dictionary: {entry}")
            return False
    
    # Additional checks based on chart type
    chart_type = chart_info['type'].lower()
    if chart_type == "pie chart" and len(chart_info['data_series']) != 1:
        logging.error("Pie chart requires exactly one data series.")
        return False
    if chart_type in ["bar chart", "line chart", "scatter plot", "histogram"] and len(chart_info['data_series']) < 1:
        logging.error(f"{chart_type.capitalize()} requires at least one data series.")
        return False
    if chart_type == "kaplan-meier curve" and len(chart_info['data_series']) == 2:
        # For Kaplan-Meier, typically need time and event data
        pass
    if chart_type == "forest plot" and len(chart_info['data_series']) >= 3:
        # Forest plots require effect sizes, confidence intervals, etc.
        pass
    
    # Check for missing or non-numeric data in data_series
    for series in chart_info['data_series']:
        for entry in chart_info['data']:
            if series not in entry or not isinstance(entry[series], (int, float, str)):
                logging.error(f"Data series '{series}' has missing or invalid data in entry: {entry}")
                return False
    
    # Additional checks for new chart types
    if chart_type == "heatmap" and len(chart_info['data_series']) < 3:
        logging.error("Heatmap requires at least three data series: x, y, and value.")
        return False
    if chart_type == "waterfall" and len(chart_info['data_series']) < 2:
        logging.error("Waterfall chart requires at least two data series: categories and values.")
        return False
    if chart_type in ["box plot", "violin plot"] and len(chart_info['data_series']) < 2:
        logging.error(f"{chart_type.capitalize()} requires at least two data series: categories and values.")
        return False

    return True

def create_chart(chart_info: Dict[str, Any]):
    """
    Creates a Matplotlib figure based on the provided chart information.

    Parameters:
    - chart_info (Dict[str, Any]): Dictionary containing chart specifications.

    Returns:
    - plt.Figure: The created Matplotlib figure.
    """
    chart_type = chart_info.get('type', '').lower()
    title = chart_info.get('title', '')
    x_label = chart_info.get('x_label', '')
    y_label = chart_info.get('y_label', '')
    data_series = chart_info.get('data_series', [])
    data = chart_info.get('data', [])

    if not data:
        raise ValueError("No data available for chart creation")

    df = pd.DataFrame(data)
    # Convert numeric columns to numbers, if possible
    for col in df.columns:
        if col != x_label and df[col].dtype == object:
            df[col] = pd.to_numeric(df[col].astype(str).str.replace('%', ''), errors='coerce')

    fig, ax = plt.subplots(figsize=(10, 6))  # Adjust figure size

    try:
        if 'bar' in chart_type:
            if x_label in df.columns:
                bars = df.plot(kind='bar', x=x_label, y=data_series, ax=ax)
                for bar in bars.containers:
                    ax.bar_label(bar, label_type='edge')
            else:
                raise ValueError(f"X-axis label '{x_label}' not found in data columns.")
        elif 'line' in chart_type:
            if x_label in df.columns:
                lines = df.plot(kind='line', x=x_label, y=data_series, ax=ax, marker='o')
                for line in lines.get_lines():
                    for x, y in zip(line.get_xdata(), line.get_ydata()):
                        ax.text(x, y, f'{y:.2f}', ha='center', va='bottom')
            else:
                raise ValueError(f"X-axis label '{x_label}' not found in data columns.")
        elif 'pie' in chart_type:
            if len(data_series) == 1:
                wedges, texts, autotexts = ax.pie(df[data_series[0]], labels=df[x_label], autopct='%1.1f%%', startangle=90)
                for autotext in autotexts:
                    autotext.set_color('white')
                ax.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
            else:
                raise ValueError("Pie chart requires exactly one data series.")
        elif 'scatter' in chart_type:
            if len(data_series) == 2:
                scatter = df.plot(kind='scatter', x=data_series[0], y=data_series[1], ax=ax)
                for i, txt in enumerate(df[x_label]):
                    ax.annotate(txt, (df[data_series[0]][i], df[data_series[1]][i]))
            else:
                raise ValueError("Scatter plot requires exactly two data series.")
        elif 'histogram' in chart_type:
            if len(data_series) == 1:
                hist = df[data_series[0]].plot(kind='hist', ax=ax, bins=10, edgecolor='black')
                for patch in hist.patches:
                    ax.text(patch.get_x() + patch.get_width() / 2, patch.get_height(), f'{patch.get_height()}', ha='center', va='bottom')
            else:
                raise ValueError("Histogram requires exactly one data series.")
        elif 'kaplan-meier curve' in chart_type:
            if len(data_series) >= 2:
                kmf = KaplanMeierFitter()
                # Assume first series is duration and second is event observed
                duration_col = data_series[0]
                event_col = data_series[1]
                if len(data_series) > 2:
                    groups = data_series[2]
                    for group in df[groups].unique():
                        mask = df[groups] == group
                        kmf.fit(df[duration_col][mask], event_observed=df[event_col][mask], label=str(group))
                        kmf.plot_survival_function(ax=ax)
                else:
                    kmf.fit(df[duration_col], event_observed=df[event_col], label="All")
                    kmf.plot_survival_function(ax=ax)
                ax.set_xlabel(x_label, fontsize=12)
                ax.set_ylabel(y_label, fontsize=12)
            else:
                raise ValueError("Kaplan-Meier curve requires at least two data series: duration and event observed.")
        elif 'heatmap' in chart_type:
            if len(data_series) >= 2:
                pivot_df = df.pivot(index=data_series[0], columns=data_series[1], values=data_series[2])
                sns.heatmap(pivot_df, annot=True, cmap='YlOrRd', ax=ax)
                ax.set_xlabel(data_series[1], fontsize=12)
                ax.set_ylabel(data_series[0], fontsize=12)
            else:
                raise ValueError("Heatmap requires at least three data series: x, y, and value.")
        elif 'waterfall' in chart_type:
            if len(data_series) >= 2:
                df = df.sort_values(by=data_series[1], ascending=False)
                df['cumulative'] = df[data_series[1]].cumsum()
                df['base'] = df['cumulative'] - df[data_series[1]]
                
                colors = ['g' if x >= 0 else 'r' for x in df[data_series[1]]]
                ax.bar(df[data_series[0]], df[data_series[1]], bottom=df['base'], color=colors)
                
                for i, (index, row) in enumerate(df.iterrows()):
                    ax.text(i, row['cumulative'], f'{row[data_series[1]]:.1f}', 
                            ha='center', va='bottom' if row[data_series[1]] >= 0 else 'top')
                
                ax.set_xlabel(data_series[0], fontsize=12)
                ax.set_ylabel(data_series[1], fontsize=12)
                ax.set_xticklabels(df[data_series[0]], rotation=45, ha='right')
            else:
                raise ValueError("Waterfall chart requires at least two data series: categories and values.")
        elif 'box plot' in chart_type:
            if len(data_series) >= 2:
                sns.boxplot(x=data_series[0], y=data_series[1], data=df, ax=ax)
                ax.set_xlabel(data_series[0], fontsize=12)
                ax.set_ylabel(data_series[1], fontsize=12)
            else:
                raise ValueError("Box plot requires at least two data series: categories and values.")
        elif 'violin plot' in chart_type:
            if len(data_series) >= 2:
                sns.violinplot(x=data_series[0], y=data_series[1], data=df, ax=ax)
                ax.set_xlabel(data_series[0], fontsize=12)
                ax.set_ylabel(data_series[1], fontsize=12)
            else:
                raise ValueError("Violin plot requires at least two data series: categories and values.")
        else:
            raise ValueError(f"Unsupported chart type: {chart_type}")

        ax.set_title(title, fontsize=14)
        plt.tight_layout()

        return fig

    except Exception as e:
        plt.close(fig)
        logging.error(f"Error creating chart '{title}': {str(e)}")
        raise

def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def extract_text_from_docx(file):
    doc = docx.Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def extract_text_from_txt(file):
    """
    Extracts text from a TXT file.

    Parameters:
    - file: The uploaded TXT file.

    Returns:
    - str: The extracted text.
    """
    return file.read().decode('utf-8')

def extract_text_from_excel(file):
    """
    Extracts text from an Excel file (XLS or XLSX).

    Parameters:
    - file: The uploaded Excel file.

    Returns:
    - str: The extracted text concatenated from all sheets.
    """
    df = pd.read_excel(file, sheet_name=None)  # Read all sheets
    text = ""
    for sheet_name, sheet_data in df.items():
        text += f"### Sheet: {sheet_name} ###\n\n"
        text += sheet_data.to_csv(index=False)
        text += "\n\n"
    return text

def extract_text_from_csv(file):
    """
    Extracts text from a CSV file.

    Parameters:
    - file: The uploaded CSV file.

    Returns:
    - str: The extracted text from the CSV.
    """
    df = pd.read_csv(file)
    text = df.to_csv(index=False)
    return text

def combine_uploaded_files(files) -> str:
    """
    Combines text extracted from multiple uploaded files.

    Parameters:
    - files: List of uploaded files.

    Returns:
    - str: Combined text from all files.
    """
    combined_text = ""
    for uploaded_file in files:
        st.write(f"Processing file: {uploaded_file.name}")
        if uploaded_file.type == "application/pdf":
            text = extract_text_from_pdf(uploaded_file)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = extract_text_from_docx(uploaded_file)
        elif uploaded_file.type == "text/plain":
            text = extract_text_from_txt(uploaded_file)
        elif uploaded_file.type in ["application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]:
            text = extract_text_from_excel(uploaded_file)
        elif uploaded_file.type == "text/csv":
            text = extract_text_from_csv(uploaded_file)
        else:
            st.warning(f"Unsupported file type: {uploaded_file.type}")
            continue
        combined_text += f"\n\n### {uploaded_file.name} ###\n\n{text}"
    return combined_text

import re
from textwrap import wrap

def generate_word_document(content: str, charts: List[Dict[str, Any]], output_format: str = "word") -> BytesIO:
    """
    Generates a Word or PDF document from the generated content and charts.

    Parameters:
    - content (str): The generated document content in Markdown.
    - charts (List[Dict[str, Any]]): List of chart information dictionaries.
    - output_format (str): 'word' or 'pdf'.

    Returns:
    - BytesIO: The generated document as a BytesIO object.
    """
    if output_format == "word":
        doc = Document()
        # Remove the '## Visualizations' section from content
        content_without_visualizations = re.sub(r'##\s+Visualizations\s*[\s\S]*', '', content, flags=re.IGNORECASE)
        # Convert Markdown to HTML
        html_content = markdown2.markdown(content_without_visualizations)
        # Parse HTML content and add to Word document
        for line in html_content.split('\n'):
            if line.startswith("<h2>"):
                para = doc.add_heading(line.replace("<h2>", "").replace("</h2>", ""), level=2)
            elif line.startswith("<h3>"):
                para = doc.add_heading(line.replace("<h3>", "").replace("</h3>", ""), level=3)
            elif line.startswith("<p>"):
                para = doc.add_paragraph(line.replace("<p>", "").replace("</p>", ""))
            elif line.startswith("<ul>"):
                para = doc.add_paragraph()
                para.style = 'List Bullet'
                for item in re.findall(r'<li>(.*?)</li>', line):
                    para.add_run(item)
            elif line.startswith("<ol>"):
                para = doc.add_paragraph()
                para.style = 'List Number'
                for item in re.findall(r'<li>(.*?)</li>', line):
                    para.add_run(item)
            elif line.startswith("<table>"):
                table_html = line + "\n"
                while not line.startswith("</table>"):
                    line = next(html_content.split('\n'))
                    table_html += line + "\n"
                table = doc.add_table(rows=0, cols=0)
                for row in re.findall(r'<tr>(.*?)</tr>', table_html):
                    cells = re.findall(r'<t[dh]>(.*?)</t[dh]>', row)
                    if len(cells) > 0:
                        row_cells = table.add_row().cells
                        for i, cell in enumerate(cells):
                            row_cells[i].text = cell
                table.style = 'Table Grid'
                for cell in table.columns[0].cells:
                    cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w'))))
            else:
                para = doc.add_paragraph(line)
        # Add charts
        if charts:
            doc.add_heading("Visualizations", level=2)
            for chart in charts:
                # Create chart using Matplotlib
                fig = create_chart(chart)
                with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmpfile:
                    fig.savefig(tmpfile.name, format='png', bbox_inches='tight')
                    doc.add_picture(tmpfile.name, width=Inches(6))
                os.unlink(tmpfile.name)
        # Save to BytesIO
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
    elif output_format == "pdf":
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter,
                                rightMargin=inch, leftMargin=inch,
                                topMargin=inch, bottomMargin=inch)
        
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='Justify', alignment=TA_JUSTIFY))
        
        elements = []

        # Remove the '## Visualizations' section from content
        content_without_visualizations = re.sub(r'##\s+Visualizations\s*[\s\S]*', '', content, flags=re.IGNORECASE)

        # Convert Markdown to HTML
        html_content = markdown2.markdown(content_without_visualizations)

        for line in html_content.split('\n'):
            if line.startswith("<h2>"):
                elements.append(Paragraph(line.replace("<h2>", "").replace("</h2>", ""), styles['Heading2']))
                elements.append(Spacer(1, 12))
            elif line.startswith("<h3>"):
                elements.append(Paragraph(line.replace("<h3>", "").replace("</h3>", ""), styles['Heading3']))
                elements.append(Spacer(1, 6))
            elif line.startswith("<p>"):
                elements.append(Paragraph(line.replace("<p>", "").replace("</p>", ""), styles['Justify']))
                elements.append(Spacer(1, 6))
            elif line.startswith("<ul>"):
                for item in re.findall(r'<li>(.*?)</li>', line):
                    elements.append(Paragraph(f"• {item}", styles['Justify']))
                    elements.append(Spacer(1, 6))
            elif line.startswith("<ol>"):
                for i, item in enumerate(re.findall(r'<li>(.*?)</li>', line), 1):
                    elements.append(Paragraph(f"{i}. {item}", styles['Justify']))
                    elements.append(Spacer(1, 6))
            elif line.startswith("<table>"):
                table_html = line + "\n"
                while not line.startswith("</table>"):
                    line = next(html_content.split('\n'))
                    table_html += line + "\n"
                data = []
                for row in re.findall(r'<tr>(.*?)</tr>', table_html):
                    cells = re.findall(r'<t[dh]>(.*?)</t[dh]>', row)
                    if len(cells) > 0:
                        data.append(cells)
                table = Table(data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                elements.append(table)
                elements.append(Spacer(1, 12))
            else:
                elements.append(Paragraph(line, styles['Justify']))
                elements.append(Spacer(1, 6))

        # Add charts
        if charts:
            elements.append(Paragraph("Visualizations", styles['Heading2']))
            elements.append(Spacer(1, 12))
            for chart in charts:
                fig = create_chart(chart)
                img_buffer = BytesIO()
                fig.savefig(img_buffer, format='png', bbox_inches='tight', dpi=300)
                img_buffer.seek(0)
                img = Image(img_buffer)
                
                # Calculate aspect ratio and adjust image size
                fig_width, fig_height = fig.get_size_inches()
                aspect_ratio = fig_height / fig_width
                
                img_width = 6 * inch  # Set a maximum width
                img_height = img_width * aspect_ratio
                
                # If the height is too large, adjust both width and height
                if img_height > 8 * inch:
                    img_height = 8 * inch
                    img_width = img_height / aspect_ratio
                
                img.drawHeight = img_height
                img.drawWidth = img_width
                img.hAlign = 'CENTER'
                
                # Add more space before the chart
                elements.append(Spacer(1, 24))
                elements.append(img)
                # Add more space after the chart
                elements.append(Spacer(1, 24))
                
                # Add chart title
                if 'title' in chart:
                    elements.append(Paragraph(chart['title'], styles['Heading4']))
                    elements.append(Spacer(1, 12))
                
                plt.close(fig)

        doc.build(elements)
        buffer.seek(0)
        return buffer
    else:
        raise ValueError("Unsupported output format. Choose 'word' or 'pdf'.")

# Add this dictionary to store the recommendations for analysis types
ANALYSIS_SOURCE_RECOMMENDATIONS = {
    "Primary Efficacy Analysis": [
        "Clinical Study Report (CSR)",
        "Study Protocol",
        "Statistical Analysis Plan (SAP)",
        "Raw Efficacy Data",
        "Tables, Listings, and Figures (TLFs) for Primary Endpoints"
    ],
    "Safety Analysis": [
        "Clinical Study Report (CSR)",
        "Safety Data (e.g., Adverse Events, Lab Data)",
        "Study Protocol",
        "Safety Monitoring Plan",
        "Tables, Listings, and Figures (TLFs) for Safety Parameters"
    ],
    "Pharmacokinetic (PK) Analysis": [
        "PK Data Set",
        "Bioanalytical Report",
        "PK Analysis Plan",
        "Study Protocol",
        "PK Modeling Results"
    ],
    "Subgroup Analysis": [
        "Clinical Study Report (CSR)",
        "Statistical Analysis Plan (SAP)",
        "Raw Data for Relevant Subgroups",
        "Pre-specified Subgroup Definitions",
        "Tables, Listings, and Figures (TLFs) for Subgroup Analyses"
    ],
    "Interim Analysis": [
        "Interim Clinical Study Report",
        "Study Protocol",
        "Interim Analysis Plan",
        "Data Cut-off Specifications",
        "Independent Data Monitoring Committee (IDMC) Reports"
    ],
    "Post-hoc Analysis": [
        "Clinical Study Report (CSR)",
        "Raw Data Sets",
        "Post-hoc Analysis Plan",
        "Justification for Additional Analyses",
        "Previous Publication Manuscripts (if applicable)"
    ],
    "Trial in Progress": [
        "Study Protocol",
        "Investigator's Brochure",
        "Clinical Trial Registration Information",
        "Enrollment Data (if available)",
        "Preliminary Safety Data (if available)"
    ],
    "Baseline Characteristics": [
        "Clinical Study Report (CSR) or Interim Report",
        "Study Protocol",
        "Statistical Analysis Plan (SAP)",
        "Demographic and Baseline Data",
        "Tables, Listings, and Figures (TLFs) for Baseline Characteristics"
    ]
}

def get_source_document_recommendations(analysis_type: str) -> List[str]:
    """
    Returns a list of recommended source documents for the given analysis type.
    
    Parameters:
    - analysis_type (str): The type of analysis being performed.
    
    Returns:
    - List[str]: A list of recommended source documents.
    """
    return ANALYSIS_SOURCE_RECOMMENDATIONS.get(analysis_type, [])

def main():
    st.title("Publication Copilot")

    publication_type = st.selectbox("Select publication type", list(PUBLICATION_TYPES.keys()))
    analysis_type = st.selectbox("Select analysis type", list(ANALYSIS_TYPES.keys()))

    # Display recommendations based on analysis type
    recommendations = get_source_document_recommendations(analysis_type)
    if recommendations:
        st.subheader("Recommended Source Documents:")
        for doc in recommendations:
            st.write(f"- {doc}")
    
    uploaded_files = st.file_uploader(
        "Upload one or more PDF, Word, TXT, XLS, or CSV documents",
        type=["pdf", "docx", "txt", "xls", "xlsx", "csv"],
        accept_multiple_files=True
    )
    
    if uploaded_files:
        user_input = combine_uploaded_files(uploaded_files)
        st.success(f"{len(uploaded_files)} file(s) uploaded and text extracted successfully!")
    else:
        user_input = st.text_area("Or enter your clinical study information:", height=300)

    additional_instructions = st.text_area(
        "Additional instructions (optional):",
        height=100,
        help="You can provide specific instructions or preferences here. For example:\n"
             "- Emphasize certain aspects of the study\n"
             "- Request specific statistical analyses\n"
             "- Ask for a particular writing style or tone\n"
             "- Specify any content that should be excluded\n"
             "- Request focus on certain subgroups or outcomes",
        placeholder="E.g., 'Please emphasize the safety profile of the drug.' or 'Focus on the subgroup analysis for patients over 65.'"
    )

    # Select output format
    output_format = st.selectbox(
        "Select output format",
        ["Word Document", "PDF"],
        help="Choose the format for the generated publication."
    )

    if st.button("Generate"):
        if user_input.strip():
            with st.spinner("Generating content..."):
                try:
                    result = generate_document_cached(publication_type, analysis_type, user_input, additional_instructions)

                    if result:
                        if result["content"].startswith("An error occurred"):
                            st.error(result["content"])
                        else:
                            # Display the generated content
                            st.subheader("Generated Content:")
                            content_without_visualizations = re.sub(r'##\s+Visualizations\s*[\s\S]*', '', result["content"], flags=re.IGNORECASE)
                            st.markdown(content_without_visualizations, unsafe_allow_html=True)

                            # Extract charts from the 'Visualizations' section
                            charts = extract_chart_info(result["content"])

                            if charts:
                                st.subheader("Visualizations:")
                                for chart_info in charts:
                                    if validate_chart_data(chart_info):
                                        try:
                                            fig = create_chart(chart_info)
                                            st.pyplot(fig)
                                        except Exception as e:
                                            st.warning(f"Could not create chart '{chart_info.get('title', 'Untitled')}': {str(e)}. Please check the chart data.")
                                            logging.error(f"Error creating chart '{chart_info.get('title', 'Untitled')}': {str(e)}")
                                            st.write("Chart data:")
                                            st.json(chart_info)
                                    else:
                                        st.warning("Received invalid chart data. Unable to visualize this chart.")
                            else:
                                st.info("No charts were generated for this content.")

                            # Assess content quality
                            quality_assessment = assess_content_quality(result["content"], publication_type, analysis_type)

                            # Display user-friendly quality assessment
                            st.subheader("Content Quality Assessment:")
                            
                            # Total word count
                            st.write(f"Total Words: {quality_assessment['total_words']}")
                            
                            # Readability score
                            fk_grade = quality_assessment['readability']['flesch_kincaid_grade']
                            if publication_type == "Plain Language Summary":
                                if 6 <= fk_grade <= 8:
                                    readability = "Excellent"
                                elif 5 <= fk_grade < 6 or 8 < fk_grade <= 9:
                                    readability = "Good"
                                elif 4 <= fk_grade < 5 or 9 < fk_grade <= 10:
                                    readability = "Fair"
                                else:
                                    readability = "Needs Improvement"
                                st.write(f"Readability: {readability} (Flesch-Kincaid Grade Level: {fk_grade:.1f})")
                                st.write("Note: For Plain Language Summaries, aim for a 6th to 8th-grade reading level.")
                            else:
                                if fk_grade < 10:
                                    readability = "Excellent"
                                elif fk_grade < 12:
                                    readability = "Good"
                                elif fk_grade < 14:
                                    readability = "Fair"
                                else:
                                    readability = "Challenging"
                                st.write(f"Readability: {readability} (Flesch-Kincaid Grade Level: {fk_grade:.1f})")
                            
                            # Section balance
                            st.write("Section Balance:")
                            total_words = sum(quality_assessment['word_counts'].values())
                            for section, count in quality_assessment['word_counts'].items():
                                percentage = (count / total_words) * 100
                                st.write(f"- {section}: {count} words ({percentage:.1f}%)")
                            
                            # Top keywords
                            st.write("Top Keywords:")
                            for word, density in list(quality_assessment['keyword_density'].items())[:5]:
                                st.write(f"- {word}: {density:.2%}")
                            
                            # Citation count
                            citation_count = quality_assessment['citation_count']
                            if citation_count == 0:
                                citation_assessment = "No citations found. Consider adding relevant citations to support your arguments."
                            elif citation_count < 5:
                                citation_assessment = "Few citations found. Consider adding more to strengthen your arguments."
                            else:
                                citation_assessment = f"Good number of citations ({citation_count})."
                            st.write(f"Citations: {citation_assessment}")
                            
                            # AI Evaluation
                            st.write("AI Evaluation:")
                            st.write(quality_assessment['ai_evaluation'])

                            # Generate downloadable document
                            with st.spinner("Generating downloadable document..."):
                                try:
                                    selected_format = "word" if output_format == "Word Document" else "pdf"
                                    document = generate_word_document(result["content"], charts, output_format=selected_format)
                                    file_extension = "docx" if selected_format == "word" else "pdf"
                                    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if selected_format == "word" else "application/pdf"
                                    st.download_button(
                                        label=f"Download as {output_format}",
                                        data=document,
                                        file_name=f"{publication_type.lower().replace(' ', '_')}_{analysis_type.lower().replace(' ', '_')}.{file_extension}",
                                        mime=mime_type
                                    )
                                except Exception as e:
                                    st.error(f"Error generating downloadable document: {str(e)}")
                            
                            # Optionally, allow downloading raw content
                            st.download_button(
                                label="Download Raw Content as Text",
                                data=result["content"],
                                file_name=f"{publication_type.lower().replace(' ', '_')}_{analysis_type.lower().replace(' ', '_')}.txt",
                                mime="text/plain"
                            )
                    else:
                        st.warning("No content was generated. Please try again.")
                except Exception as e:
                    st.error(f"An unexpected error occurred: {str(e)}")
                    logging.exception("An unexpected error occurred in the main application:")
        else:
            st.warning("Please enter some information or upload at least one file before generating.")

if __name__ == '__main__':
    logging.debug("Entering main block")
    try:
        main()
    except Exception as e:
        logging.exception("An error occurred in the main application:")
        st.error(f"An error occurred: {str(e)}")
